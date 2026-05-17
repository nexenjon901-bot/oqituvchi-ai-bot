import os
import logging
import asyncio
import sqlite3
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import CommandStart
from aiogram.utils.markdown import hbold
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, FSInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from google import genai
from google.genai import types as ai_types
from dotenv import load_dotenv

# PDF va Word yaratish uchun kutubxonalar
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document

load_dotenv()

TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
AI_API_KEY = os.getenv("GEMINI_API_KEY")
ADMIN_ID = 8407085035  # Bu yerga o'zingizning Telegram ID raqamingizni yozing

if not TOKEN or not AI_API_KEY:
    raise ValueError("XATOLIK: .env faylida tokenlar to'liq emas!")

client = genai.Client(api_key=AI_API_KEY)
dp = Dispatcher()

# Ma'lumotlar bazasini sozlash
conn = sqlite3.connect("users_v3.db", check_same_thread=False)
cursor = conn.cursor()
cursor.execute("""
    CREATE TABLE IF NOT EXISTS users (
        user_id INTEGER PRIMARY KEY,
        username TEXT,
        full_name TEXT
    )
""")
cursor.execute("""
    CREATE TABLE IF NOT EXISTS history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        role TEXT,
        text TEXT
    )
""")
conn.commit()

class BotStates(StatesGroup):
    waiting_for_ad = State()

SYSTEM_PROMPT = "SEN “O‘QITUVCHI AI” NOMLI PROFESSIONAL SUN’IY INTELLEKT YORDAMCHISISAN. BERILGAN SAVOLLARGA ANIQ VA TUSHUNARLI JAVOB BER."

# Asosiy menyu tugmalari
def get_main_keyboard(user_id):
    buttons = [
        [KeyboardButton(text="📝 Test / Quiz yaratish"), KeyboardButton(text="📚 Referat / Insho yozish")],
        [KeyboardButton(text="🧮 Matematika & Masalalar"), KeyboardButton(text="💡 AI Ustozdan so'rash")],
        [KeyboardButton(text="📄 PDF yuklab olish"), KeyboardButton(text="📝 Word (Docx) yuklab olish")],
        [KeyboardButton(text="🔄 Tarixni tozalash")]
    ]
    if user_id == ADMIN_ID:
        buttons.append([KeyboardButton(text="⚙️ Admin Panel")])
    return ReplyKeyboardMarkup(keyboard=buttons, resize_keyboard=True)

# Admin menyusi
def get_admin_keyboard():
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📊 Foydalanuvchilar statistikasi")],
            [KeyboardButton(text="📢 Hammasiga xabar yuborish")],
            [KeyboardButton(text="⬅️ Bosh menyuga qaytish")]
        ],
        resize_keyboard=True
    )

# PDF yaratish funksiyasi
def create_pdf(text, filename="Oqituvchi_AI.pdf"):
    doc = SimpleDocTemplate(filename, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('PDFTitle', parent=styles['Heading1'], fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=20, textColor="#1A237E")
    body_style = ParagraphStyle('PDFBody', parent=styles['Normal'], fontSize=12, leading=18, alignment=TA_JUSTIFY, spaceAfter=10)
    story = [Paragraph("<b>“O‘QITUVCHI AI” RESURSI</b>", title_style), Spacer(1, 15)]
    for line in text.split("\n"):
        if line.strip():
            clean_line = line.replace("**", "<b>").replace("__", "<i>")
            story.append(Paragraph(clean_line, body_style))
    doc.build(story)
    return filename

# Word yaratish funksiyasi
def create_word(text, filename="Oqituvchi_AI.docx"):
    doc = Document()
    doc.add_heading("“O‘QITUVCHI AI” TAQDIM ETADI", level=1)
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.replace("**", "").replace("__", ""))
    doc.save(filename)
    return filename

# Start buyrug'i
@dp.message(CommandStart())
async def command_start_handler(message: types.Message) -> None:
    user_id = message.from_user.id
    cursor.execute("INSERT OR IGNORE INTO users (user_id, username, full_name) VALUES (?, ?, ?)", 
                   (user_id, message.from_user.username, message.from_user.full_name))
    conn.commit()
    
    welcome_text = (
        f"Assalomu alaykum, hurmatli {message.from_user.full_name}!\n\n"
        f"Men — O‘QITUVCHI AI professional sun'iy intellekt yordamchisiman.\n\n"
        f"Men sizga quyidagi ishlarda yaqindan ko'maklasha olaman:\n"
        f"• Istalgan mavzuda mukammal Test va Quizlar yaratish\n"
        f"• Sifatli Referat, Insho va Esse yozish\n"
        f"• Matematika, Mantiq va Buxgalteriya masalalarini yechish\n"
        f"• Ingliz va Rus tillarini o'rganish va matnlarni tarjima qilish\n"
        f"• Tayyor ma'lumotlarni PDF yoki Word (Docx) formatida yuklab olish\n\n"
        f"Hatto menga daftaringizdagi qiyin misollarni rasmga olib tashlasangiz ham yechib bera olaman!\n\n"
        f"Boshlash uchun pastdagi menyudan o'zingizga kerakli bo'limni tanlang:"
    )
    await message.answer(welcome_text, reply_markup=get_main_keyboard(user_id))

# Tarixni tozalash
@dp.message(F.text == "🔄 Tarixni tozalash")
async def clear_history(message: types.Message):
    cursor.execute("DELETE FROM history WHERE user_id = ?", (message.from_user.id,))
    conn.commit()
    await message.answer("Suhbatimiz tarixi tozalandi! Endi yangi mavzuda gaplashishimiz mumkin.")

# Rasm qabul qilish
@dp.message(F.photo)
async def photo_handler(message: types.Message):
    user_id = message.from_user.id
    await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
    
    photo = message.photo[-1]
    file_info = await message.bot.get_file(photo.file_id)
    destination = f"photo_{user_id}.jpg"
    await message.bot.download_file(file_info.file_path, destination)
    
    caption = message.caption if message.caption else "Ushbu rasmdagi topshiriq yoki misolni yechib, tushuntirib ber."
    await message.answer("Rasm qabul qilindi. AI tahlil qilmoqda, iltimos kuting...")
    
    try:
        with open(destination, "rb") as f:
            image_bytes = f.read()
            
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                ai_types.Part.from_bytes(data=image_bytes, mime_type="image/jpeg"),
                caption
            ],
            config=ai_types.GenerateContentConfig(system_instruction=SYSTEM_PROMPT)
        )
        ai_reply = response.text
        
        cursor.execute("INSERT INTO history (user_id, role, text) VALUES (?, ?, ?)", (user_id, 'user', f"[Rasm]: {caption}"))
        cursor.execute("INSERT INTO history (user_id, role, text) VALUES (?, ?, ?)", (user_id, 'model', ai_reply))
        conn.commit()
        
        await message.answer(ai_reply)
    except Exception as e:
        await message.answer(f"Xatolik yuz berdi: {e}")
    finally:
        if os.path.exists(destination):
            os.remove(destination)

# Umumiy xabarlar va Admin panel boshqaruvi
@dp.message()
async def main_handler(message: types.Message, state: FSMContext) -> None:
    user_text = message.text
    user_id = message.from_user.id
    
    if user_text == "⚙️ Admin Panel" and user_id == ADMIN_ID:
        await message.answer("Admin panel:", reply_markup=get_admin_keyboard())
        return
    elif user_text == "📊 Foydalanuvchilar statistikasi" and user_id == ADMIN_ID:
        cursor.execute("SELECT COUNT(*) FROM users")
        await message.answer(f"Jami foydalanuvchilar: {cursor.fetchone()[0]} ta")
        return
    elif user_text == "📢 Hammasiga xabar yuborish" and user_id == ADMIN_ID:
        await message.answer("Foydalanuvchilarga yubormoqchi bo'lgan xabarni yozing:")
        await state.set_state(BotStates.waiting_for_ad)
        return
    elif user_text == "⬅️ Bosh menyuga qaytish":
        await message.answer("Bosh menyu:", reply_markup=get_main_keyboard(user_id))
        return

    current_state = await state.get_state()
    if current_state == BotStates.waiting_for_ad.state and user_id == ADMIN_ID:
        cursor.execute("SELECT user_id FROM users")
        users = cursor.fetchall()
        for u in users:
            try: 
                await message.copy_to(chat_id=u[0])
                await asyncio.sleep(0.05)
            except: 
                pass
        await message.answer("Xabar barcha foydalanuvchilarga tarqatildi.", reply_markup=get_admin_keyboard())
        await state.clear()
        return

    if user_text in ["📄 PDF yuklab olish", "📝 Word (Docx) yuklab olish"]:
        cursor.execute("SELECT text FROM history WHERE user_id = ? AND role = 'model' ORDER BY id DESC LIMIT 1", (user_id,))
        last_row = cursor.fetchone()
        if not last_row:
            await message.answer("Avval biron bir mavzuda ma'lumot so'rang.")
            return
            
        last_response = last_row[0]
        if "PDF" in user_text:
            file_path = create_pdf(last_response)
            await message.answer_document(document=FSInputFile(file_path), caption="Tayyor PDF hujjat.")
        else:
            file_path = create_word(last_response)
            await message.answer_document(document=FSInputFile(file_path), caption="Tayyor Word hujjat.")
        os.remove(file_path)
        return

    if user_text in ["📝 Test / Quiz yaratish", "📚 Referat / Insho yozish", "🧮 Matematika & Masalalar", "💡 AI Ustozdan so'rash"]:
        await message.answer(f"Siz '{user_text}' bo'limini tanladingiz. Mavzu yoki savolingizni batafsil yozib yuboring:")
        return

    # Sun'iy intellekt javobi va xotira (kontekst) qismi
    await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")
    
    cursor.execute("SELECT role, text FROM history WHERE user_id = ? ORDER BY id DESC LIMIT 6", (user_id,))
    rows = cursor.fetchall()[::-1]
    
    contents = []
    for r in rows:
        contents.append(ai_types.Content(role=r[0], parts=[ai_types.Part.from_text(text=r[1])]))
    contents.append(ai_types.Content(role="user", parts=[ai_types.Part.from_text(text=user_text)]))
    
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=contents,
            config=ai_types.GenerateContentConfig(system_instruction=SYSTEM_PROMPT),
        )
        ai_reply = response.text
        
        cursor.execute("INSERT INTO history (user_id, role, text) VALUES (?, ?, ?)", (user_id, 'user', user_text))
        cursor.execute("INSERT INTO history (user_id, role, text) VALUES (?, ?, ?)", (user_id, 'model', ai_reply))
        conn.commit()
        
        await message.answer(ai_reply)
    except Exception as e:
        logging.error(f"Xatolik: {e}")
        await message.answer("Kechirasiz, xatolik yuz berdi. Qayta urinib ko'ring.")

async def main() -> None:
    bot = Bot(token=TOKEN)
    await dp.start_polling(bot)

if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO, stream=sys.stdout)
    try: 
        asyncio.run(main())
    except KeyboardInterrupt: 
        print("Bot to'xtatildi")
